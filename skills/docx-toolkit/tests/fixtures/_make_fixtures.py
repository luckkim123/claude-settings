"""Generate test fixtures for docx_inspect.

Run from this directory:
    cd tests/fixtures && ../../.venv/bin/python _make_fixtures.py
"""
from docx import Document


def make_plain_body():
    doc = Document()
    doc.add_paragraph("hello world")
    doc.add_paragraph("second paragraph")
    doc.save("plain_body.docx")


def make_with_header_styles():
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "Header text"
    # Add 12 user-defined paragraph styles (>10 triggers high-risk classification)
    for i in range(12):
        doc.styles.add_style(f"CustomStyle{i}", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    doc.add_paragraph("body content")
    doc.save("with_header_styles.docx")


def make_with_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save("with_table.docx")


if __name__ == "__main__":
    make_plain_body()
    make_with_header_styles()
    make_with_table()
    print("Generated: plain_body.docx, with_header_styles.docx, with_table.docx")
