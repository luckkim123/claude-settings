#!/usr/bin/env python3
"""Inspect a .docx and classify edit risk for routing decisions.

Usage:
    docx_inspect.py <path>
    docx_inspect.py <path> --compare <other>

Outputs JSON to stdout:
    {
      "path": "...",
      "signals": { has_header, has_footer, user_style_count,
                   has_table, has_image, has_theme },
      "risk": "low" | "medium" | "high",
      "compare": { "<member>": { "a": "<sha>", "b": "<sha>", "match": bool } }  # if --compare
    }
"""
import hashlib
import json
import sys
import zipfile
from pathlib import Path

from docx import Document

# python-docx ships these out of the box; user-defined styles are anything else.
_BUILTIN_STYLE_NAMES = {
    "Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5",
    "Heading 6", "Heading 7", "Heading 8", "Heading 9", "Title", "Subtitle",
    "List Paragraph", "Quote", "Intense Quote", "Caption", "TOC Heading",
    "Default Paragraph Font", "No List", "Table Normal", "Header", "Footer",
    "Body Text", "Body Text 2", "Body Text 3", "List Bullet", "List Number",
    "List Continue", "Macro Text", "TOC 1", "TOC 2", "TOC 3", "TOC 4",
    "TOC 5", "TOC 6", "TOC 7", "TOC 8", "TOC 9", "Hyperlink", "FollowedHyperlink",
}

_COMPARE_MEMBERS = (
    "word/document.xml",
    "word/styles.xml",
    "word/theme/theme1.xml",
    "word/header1.xml",
    "word/footer1.xml",
    "word/numbering.xml",
)


def _signals(path: Path) -> dict:
    doc = Document(str(path))
    has_header = any(
        any(p.text.strip() for p in s.header.paragraphs)
        for s in doc.sections
    )
    has_footer = any(
        any(p.text.strip() for p in s.footer.paragraphs)
        for s in doc.sections
    )
    user_styles = [s for s in doc.styles if s.name not in _BUILTIN_STYLE_NAMES]
    has_table = len(doc.tables) > 0
    has_image = any(
        "graphicData" in r._element.xml
        for p in doc.paragraphs
        for r in p.runs
    )
    with zipfile.ZipFile(path) as z:
        has_theme = "word/theme/theme1.xml" in z.namelist()
    return {
        "has_header": has_header,
        "has_footer": has_footer,
        "user_style_count": len(user_styles),
        "has_table": has_table,
        "has_image": has_image,
        "has_theme": has_theme,
    }


def _classify(s: dict) -> str:
    if s["has_header"] or s["has_footer"]:
        return "high"
    if s["user_style_count"] > 10:
        return "high"
    if s["has_image"]:
        return "high"
    if s["has_table"]:
        return "medium"
    return "low"


def _hash_member(path: Path, member: str) -> str | None:
    with zipfile.ZipFile(path) as z:
        if member not in z.namelist():
            return None
        return hashlib.sha256(z.read(member)).hexdigest()


def _compare(a: Path, b: Path) -> dict:
    out = {}
    for m in _COMPARE_MEMBERS:
        ha = _hash_member(a, m)
        hb = _hash_member(b, m)
        out[m] = {"a": ha, "b": hb, "match": ha == hb}
    return out


def main(argv: list[str]) -> int:
    if not argv:
        print("usage: docx_inspect.py <path> [--compare <other>]", file=sys.stderr)
        return 2

    path = Path(argv[0])
    if not path.is_file():
        print(f"file not found: {path}", file=sys.stderr)
        return 2

    out = {"path": str(path), "signals": _signals(path)}
    out["risk"] = _classify(out["signals"])

    if "--compare" in argv:
        idx = argv.index("--compare")
        if idx + 1 >= len(argv):
            print("--compare requires a path argument", file=sys.stderr)
            return 2
        other = Path(argv[idx + 1])
        if not other.is_file():
            print(f"compare file not found: {other}", file=sys.stderr)
            return 2
        out["compare"] = _compare(path, other)

    json.dump(out, sys.stdout)
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
